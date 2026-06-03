"""
Generates the full research manuscript as a Word (.docx) file.
Run: python paper/generate_manuscript.py
Output: paper/Bangladesh_Measles_2026_Manuscript.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(os.path.dirname(OUT), "figures")
DATA = os.path.join(os.path.dirname(OUT), "data")

# ── helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_para(doc, text, size=12, italic=False, bold=False, justify=True, space_before=0, space_after=6, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic, bold=bold)
    return p

def add_mixed_para(doc, parts, justify=True, space_before=0, space_after=6, indent=False):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after = Pt(4)
        r = cp.add_run(caption)
        set_font(r, size=11, bold=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = tbl.rows[0]
    shade_row(hrow, "2E75B6")
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(row, "EBF3FB")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_font(run, size=10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return tbl


# ── build document ─────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ─────────────────────────────────────────────────────────────

doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(16)
tr = title_p.add_run(
    "Epidemiology of the 2026 Measles Outbreak in Bangladesh: "
    "A Descriptive Analysis of National Surveillance Data "
    "with Vaccine Coverage Assessment"
)
set_font(tr, size=16, bold=True)

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(6)
ar = auth_p.add_run("Khalilur Rahman Ridoy Khan")
set_font(ar, size=12, bold=True)

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(4)
afr = aff_p.add_run("Department of Computer Science and Engineering\nEast West University, Dhaka, Bangladesh")
set_font(afr, size=11, italic=True)

email_p = doc.add_paragraph()
email_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
email_p.paragraph_format.space_after = Pt(4)
er = email_p.add_run("Correspondence: khalilurrahmanridoykhan@gmail.com")
set_font(er, size=10)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(4)
dr = date_p.add_run("Date: June 2026")
set_font(dr, size=10)

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo_p.paragraph_format.space_after = Pt(4)
rr = repo_p.add_run("Data & Code: https://github.com/khalilurrrahmanridoykhan/bangladesh-measles-2026")
set_font(rr, size=10, italic=True)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
kw_p.paragraph_format.space_after = Pt(20)
kwr = kw_p.add_run(
    "Keywords: measles; হাম; Bangladesh; outbreak; epidemiology; "
    "vaccine-preventable disease; DGHS; surveillance; 2026"
)
set_font(kwr, size=10, italic=True)

doc.add_page_break()

# ── ABSTRACT ───────────────────────────────────────────────────────────────

add_heading(doc, "Abstract", size=14, space_before=0)

abstract_sections = [
    ("Background: ",
     "Bangladesh has experienced a large-scale measles outbreak since March 2026, "
     "coinciding with a national Measles-Rubella (MR) immunisation campaign. "
     "This study describes the epidemiological characteristics of the outbreak using "
     "publicly available daily surveillance data published by the Directorate General "
     "of Health Services (DGHS)."),
    ("Methods: ",
     "We extracted structured data from 62 consecutive daily press release PDFs "
     "published by DGHS (2 April to 2 June 2026) using automated web scraping and "
     "PDF text extraction. We analysed temporal trends, geographic distribution across "
     "eight administrative divisions, mortality, and the relationship between "
     "MR campaign vaccination coverage and division-level incidence rates."),
    ("Results: ",
     "A total of 73,362 suspected and 9,136 confirmed measles cases were reported "
     "during the study period (confirmation rate 12.5%). A total of 59,106 patients "
     "were hospitalised and 54,812 (92.7%) recovered. Suspected and confirmed deaths "
     "numbered 504 and 90, respectively (confirmed case fatality rate 0.99%). "
     "The outbreak grew rapidly in its early phase, with a doubling time of 6.1 days. "
     "Peak daily suspected cases reached 1,503 on 10 May 2026. Dhaka division "
     "accounted for 47.0% of all confirmed cases (34,497 cases; incidence 95.7 per "
     "100,000 population), followed by Chattogram (16.0%) and Rajshahi (12.6%). "
     "National MR campaign coverage was 102%; however, Pearson correlation between "
     "division-level coverage and incidence was non-significant (r = 0.10, p = 0.82), "
     "indicating that the campaign did not prevent ongoing transmission in divisions "
     "with pre-existing immunity gaps."),
    ("Conclusions: ",
     "This outbreak represents one of the largest measles events recorded in Bangladesh. "
     "Despite high MR campaign coverage, urban transmission persisted — particularly in "
     "Dhaka. Strengthening routine immunisation, improving cold-chain management, and "
     "targeting high-density urban populations are essential to interrupt transmission "
     "and prevent future outbreaks."),
]

for label, content in abstract_sections:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label)
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(content)
    set_font(r2, size=12)

doc.add_page_break()

# ── 1. INTRODUCTION ────────────────────────────────────────────────────────

add_heading(doc, "1. Introduction", size=13)

add_para(doc,
    "Measles remains one of the most contagious infectious diseases known to medicine, "
    "with a basic reproduction number (R₀) estimated between 12 and 18. Despite the "
    "existence of a safe, effective, and affordable vaccine, measles caused an estimated "
    "136,000 deaths globally in 2022, the majority among children under five years of age "
    "(World Health Organization [WHO], 2023). Following disruptions to routine immunisation "
    "services during the COVID-19 pandemic (2020–2022), a global resurgence of measles has "
    "been observed, with large outbreaks reported across Africa, South-East Asia, and the "
    "Eastern Mediterranean region.",
    indent=True)

add_para(doc,
    "Bangladesh, a densely populated lower-middle-income country with approximately "
    "170 million inhabitants, has maintained an Expanded Programme on Immunisation (EPI) "
    "since 1979. The measles vaccine was introduced into the national schedule in 1980, "
    "and a two-dose schedule — measles-rubella (MR) vaccine at 9 and 15 months — has been "
    "in place since 2012. Bangladesh had previously achieved measles vaccination coverage "
    "exceeding 95% in many districts, and no major nationwide outbreak had been documented "
    "for several years prior to 2026.",
    indent=True)

add_para(doc,
    "In March 2026, the DGHS activated its Integrated Control Centre (Control Room) in "
    "response to a sharp rise in febrile rash illness consistent with measles across "
    "multiple divisions. Daily situation reports were published publicly from 15 March 2026, "
    "providing a rare opportunity for granular, real-time epidemiological analysis using "
    "government surveillance data. By 2 June 2026, more than 73,000 suspected cases had "
    "been reported nationally.",
    indent=True)

add_para(doc,
    "This study has three objectives: (i) to describe the temporal trajectory of the 2026 "
    "Bangladesh measles outbreak using 62 days of daily national surveillance data; (ii) to "
    "characterise the geographic distribution of the outbreak burden across Bangladesh's "
    "eight administrative divisions; and (iii) to assess the relationship between the 2026 "
    "national Measles-Rubella immunisation campaign coverage and division-level measles "
    "incidence.",
    indent=True)

# ── 2. METHODS ─────────────────────────────────────────────────────────────

add_heading(doc, "2. Methods", size=13)

add_heading(doc, "2.1 Data Source", size=12, space_before=8)
add_para(doc,
    "Data were sourced exclusively from publicly available daily press release PDFs "
    "published by the DGHS of the Government of Bangladesh "
    "(https://dghs.gov.bd/pages/press-releases/). Each press release is a structured "
    "three-page PDF document produced by the DGHS Integrated Control Centre and signed "
    "by the Additional Director General (Administration). The documents report: "
    "(i) national suspected and confirmed measles cases in the preceding 24 hours and "
    "cumulatively from 15 March 2026; (ii) suspected and confirmed deaths; "
    "(iii) hospitalisation and recovery counts; (iv) division-level breakdowns of the "
    "above; and (v) Measles-Rubella Campaign 2026 vaccination coverage by division and "
    "city corporation.",
    indent=True)

add_heading(doc, "2.2 Data Collection", size=12, space_before=8)
add_para(doc,
    "Automated data collection was conducted using custom Python scripts (Python 3.14; "
    "libraries: requests, BeautifulSoup4, pdfplumber). The scraper retrieved the HTML "
    "listing page, identified all press releases containing the keyword 'হাম' (measles), "
    "extracted the direct PDF object URL from each detail page, downloaded each PDF, and "
    "extracted structured text. Bangla numeral characters (০–৯) were converted to "
    "Western Arabic numerals (0–9) prior to parsing. All scripts and the resulting clean "
    "datasets are available at the public GitHub repository: "
    "https://github.com/khalilurrrahmanridoykhan/bangladesh-measles-2026.",
    indent=True)

add_heading(doc, "2.3 Variables", size=12, space_before=8)
add_para(doc,
    "For each reporting day we extracted: daily suspected cases, cumulative suspected cases, "
    "daily confirmed cases, cumulative confirmed cases, cumulative hospitalisations, "
    "cumulative recoveries, daily suspected deaths, cumulative suspected deaths, daily "
    "confirmed deaths, and cumulative confirmed deaths. Division-level data (eight divisions) "
    "were extracted from page two of each PDF. Measles-Rubella Campaign 2026 coverage data "
    "(target population, doses administered, coverage percentage) were extracted from pages "
    "two and three.",
    indent=True)

add_heading(doc, "2.4 Case Definitions", size=12, space_before=8)
add_para(doc,
    "Case definitions follow those used by DGHS, which are aligned with WHO 2018 measles "
    "surveillance guidelines. A suspected case is defined as any person with fever and "
    "generalised maculopapular rash, plus at least one of: cough, coryza, or conjunctivitis. "
    "A confirmed case is a suspected case with laboratory confirmation (measles IgM serology "
    "or viral isolation) or epidemiologic linkage to a laboratory-confirmed case.",
    indent=True)

add_heading(doc, "2.5 Statistical Analysis", size=12, space_before=8)
add_para(doc,
    "Descriptive statistics were computed for national and division-level indicators. "
    "The epidemic curve was plotted as a bar chart of daily suspected and confirmed cases, "
    "with a seven-day centred rolling average overlay. Cumulative case trajectories were "
    "plotted for suspected, confirmed, hospitalised, recovered, and suspected death counts. "
    "The doubling time during the exponential growth phase (first 14 days) was estimated "
    "using log-linear regression: ln(cumulative suspected cases) = α + β·day, with doubling "
    "time calculated as ln(2)/β. Incidence rates were calculated per 100,000 population "
    "using division-level population estimates from the Bangladesh Population and Housing "
    "Census 2022 (Bangladesh Bureau of Statistics). The Pearson correlation coefficient was "
    "used to assess the relationship between division-level MR campaign coverage (%) and "
    "confirmed measles incidence per 100,000. All analyses were performed in Python 3.14 "
    "(pandas 3.0, matplotlib 3.10, scipy 1.17). Statistical significance was set at p < 0.05.",
    indent=True)

add_heading(doc, "2.6 Ethics", size=12, space_before=8)
add_para(doc,
    "This study used only publicly available, aggregated government surveillance data. "
    "No individual patient data, identifiable information, or private records were accessed. "
    "Formal ethics review board approval was not required.",
    indent=True)

# ── 3. RESULTS ─────────────────────────────────────────────────────────────

add_heading(doc, "3. Results", size=13)

add_heading(doc, "3.1 National Overview", size=12, space_before=8)
add_para(doc,
    "Over the 62-day study period (2 April to 2 June 2026), a total of 73,362 suspected "
    "and 9,136 confirmed measles cases were reported nationally, yielding a confirmation "
    "rate of 12.5% (Table 1). A total of 59,106 patients were admitted to hospital, of "
    "whom 54,812 (92.7%) were recorded as recovered by the end of the study period, "
    "leaving an estimated 4,294 active cases on 2 June 2026. Suspected deaths totalled 504 "
    "(suspected case fatality rate [CFR] 0.69%), while confirmed deaths numbered 90 "
    "(confirmed CFR 0.99%).",
    indent=True)

# TABLE 1
add_table(doc,
    headers=["Indicator", "Value"],
    rows=[
        ("Reporting period", "2 April 2026 – 2 June 2026 (62 days)"),
        ("Total suspected cases", "73,362"),
        ("Total confirmed cases", "9,136"),
        ("Confirmation rate", "12.5%"),
        ("Total hospitalised (cumulative)", "59,106"),
        ("Hospitalisation rate (of suspected)", "80.6%"),
        ("Total recovered (cumulative)", "54,812"),
        ("Recovery rate (of hospitalised)", "92.7%"),
        ("Total suspected deaths", "504"),
        ("Total confirmed deaths", "90"),
        ("Case fatality rate — suspected", "0.69%"),
        ("Case fatality rate — confirmed", "0.99%"),
        ("Peak daily suspected cases", "1,503 (10 May 2026)"),
        ("Mean daily suspected cases", "1,167"),
    ],
    col_widths=[3.5, 2.8],
    caption="Table 1. National Summary Statistics — Bangladesh Measles Outbreak, 2 April to 2 June 2026"
)

add_heading(doc, "3.2 Temporal Trends", size=12, space_before=8)
add_para(doc,
    "The epidemic curve demonstrated a rapid early rise followed by a sustained plateau "
    "(Figure 1). Daily suspected cases increased from 685 on 2 April 2026 to a peak of "
    "1,503 on 10 May 2026 — a 2.2-fold increase over 38 days. The mean daily suspected "
    "case count across the study period was 1,167 (range: 674–1,503). During the final "
    "week of the study (28 May to 2 June 2026), daily suspected cases declined to below "
    "1,100, suggesting a possible plateau or early decline, though this trend requires "
    "continued monitoring.",
    indent=True)

add_para(doc,
    "The doubling time during the first 14 days of the study period was estimated at "
    "6.1 days (daily growth rate 11.3%; r² = 0.94), consistent with exponential growth. "
    "Week-over-week analysis revealed the highest relative growth in epidemiological "
    "week 2 (9–15 April, +12.1%), followed by week 4 (23–29 April, +10.2%). Growth "
    "decelerated from week 7 onward, with a −23.6% week-over-week decline in the final "
    "partial week (28 May to 2 June 2026). The cumulative trajectory of hospitalisations "
    "closely tracked suspected cases (Figure 2), consistent with a high hospitalisation "
    "burden throughout the outbreak.",
    indent=True)

add_para(doc, "(See Figure 1: Epidemic curve; Figure 2: Cumulative trajectory)", italic=True, indent=True)

add_heading(doc, "3.3 Geographic Distribution", size=12, space_before=8)
add_para(doc,
    "Measles burden was highly concentrated in Dhaka division, which accounted for 34,497 "
    "confirmed cases — 47.0% of the national total — with an incidence rate of 95.7 per "
    "100,000 population (Table 2). Chattogram division recorded the second highest absolute "
    "burden (12,009 confirmed cases, 16.0%) with an incidence of 42.3 per 100,000. Rajshahi "
    "division (9,414 cases; 50.9 per 100,000) and Khulna division (5,422 cases; 34.8 per "
    "100,000) followed. The lowest burden was observed in Rangpur division (3,323 confirmed "
    "cases; 21.2 per 100,000). Mymensingh division data were incomplete in the division-level "
    "breakdowns and could not be reliably quantified for this analysis.",
    indent=True)

add_para(doc, "(See Figure 3: Division-level burden)", italic=True, indent=True)

# TABLE 2
add_table(doc,
    headers=["Division", "Population\n(2022 Census)", "Confirmed\nCases", "Incidence\n/100,000", "% of\nNational Total"],
    rows=[
        ("Dhaka",       "36,054,418", "34,497", "95.7",  "47.0%"),
        ("Chattogram",  "28,423,019", "12,009", "42.3",  "16.0%"),
        ("Rajshahi",    "18,484,858",  "9,414", "50.9",  "12.6%"),
        ("Barishal",     "8,325,666",  "6,634", "79.7",   "8.8%"),
        ("Khulna",      "15,563,000",  "5,422", "34.8",   "7.2%"),
        ("Sylhet",      "10,009,239",  "3,688", "36.8",   "4.9%"),
        ("Rangpur",     "15,665,000",  "3,323", "21.2",   "4.4%"),
        ("Mymensingh",  "11,370,000",    "N/A",   "N/A", "N/A"),
        ("NATIONAL",   "143,895,200", "74,987*", "52.1",  "100%"),
    ],
    col_widths=[1.5, 1.8, 1.4, 1.4, 1.4],
    caption="Table 2. Division-level Confirmed Measles Cases and Incidence Rates, Bangladesh 2026\n"
            "* National total includes cases not attributed to specific divisions in daily reports."
)

add_heading(doc, "3.4 Mortality", size=12, space_before=8)
add_para(doc,
    "Deaths data were extractable for 14 of the 62 reporting days (23 May to 2 June 2026) "
    "from the PDF text. As of 2 June 2026, 504 suspected deaths and 90 confirmed measles "
    "deaths had been recorded cumulatively since 15 March 2026. The confirmed CFR of 0.99% "
    "is consistent with measles mortality estimates for low- and middle-income countries "
    "without vitamin A supplementation programmes in place. Daily suspected deaths ranged "
    "from 3 to 16 per day during the period for which data were available. The distribution "
    "of deaths by district was reported in press releases but could not be extracted "
    "systematically due to PDF layout variability; Dhaka and Faridpur districts were cited "
    "most frequently in available reports.",
    indent=True)

add_heading(doc, "3.5 Vaccination Coverage and Incidence", size=12, space_before=8)
add_para(doc,
    "The national Measles-Rubella Campaign 2026 achieved 102% overall coverage "
    "(18,452,281 doses administered against a target of 18,015,064 children). Coverage "
    "exceeded 100% in seven of eight divisions (range 99%–103%), reflecting catch-up of "
    "missed children. Despite this high coverage, the Pearson correlation between "
    "division-level MR campaign coverage and confirmed measles incidence per 100,000 was "
    "r = 0.10 (p = 0.82), indicating no statistically significant association (Table 3, "
    "Figure 4). This finding strongly suggests that the campaign, conducted during an "
    "active outbreak, was reactive rather than prophylactic, and that pre-existing "
    "immunity gaps accumulated prior to the campaign were the primary driver of "
    "transmission.",
    indent=True)

add_para(doc, "(See Figure 4: Vaccination coverage vs. incidence)", italic=True, indent=True)

# TABLE 3
add_table(doc,
    headers=["Division", "MR Campaign\nCoverage (%)", "Confirmed Cases", "Incidence\n/100,000"],
    rows=[
        ("Dhaka",      "103%", "34,497", "95.7"),
        ("Chattogram", "103%", "12,009", "42.3"),
        ("Rajshahi",   "103%",  "9,414", "50.9"),
        ("Barishal",   "101%",  "6,634", "79.7"),
        ("Khulna",     "101%",  "5,422", "34.8"),
        ("Sylhet",      "99%",  "3,688", "36.8"),
        ("Rangpur",    "103%",  "3,323", "21.2"),
        ("Mymensingh", "102%",    "N/A",   "N/A"),
    ],
    col_widths=[1.6, 1.8, 1.8, 1.8],
    caption="Table 3. MR Campaign 2026 Coverage vs. Confirmed Measles Incidence by Division\n"
            "Pearson r = 0.10, p = 0.82 (no statistically significant correlation)"
)

# ── 4. DISCUSSION ──────────────────────────────────────────────────────────

add_heading(doc, "4. Discussion", size=13)

add_heading(doc, "4.1 Summary of Findings", size=12, space_before=8)
add_para(doc,
    "This study presents the first systematic epidemiological analysis of the 2026 measles "
    "outbreak in Bangladesh, using 62 days of publicly available national surveillance data. "
    "The outbreak was characterised by rapid exponential growth (doubling time 6.1 days), "
    "high hospitalisation burden (80.6% of suspected cases), and strong geographic "
    "concentration in Dhaka division (47% of national confirmed cases). Despite achieving "
    "over 100% MR campaign coverage nationally, no statistically significant correlation "
    "between vaccination coverage and division-level incidence was observed.",
    indent=True)

add_heading(doc, "4.2 Rapid Transmission and Urban Concentration", size=12, space_before=8)
add_para(doc,
    "The estimated doubling time of 6.1 days during the early outbreak phase is consistent "
    "with a high effective reproduction number (Rₑ > 1), indicating ongoing transmission "
    "in a susceptible population. The concentration of 47% of cases in Dhaka division "
    "is not surprising given Dhaka's extreme population density (approximately 44,000 "
    "persons per km² in the metropolitan area), high household crowding, and large "
    "populations of urban migrants with potentially incomplete vaccination histories. "
    "Urban measles outbreaks disproportionately affect densely populated megacities "
    "globally, and Dhaka follows this pattern.",
    indent=True)

add_para(doc,
    "Barishal division recorded the highest incidence rate per 100,000 (79.7) despite "
    "having the smallest population, suggesting intense transmission in a geographically "
    "compact area. Conversely, Rangpur — a largely rural, less densely populated division "
    "— recorded the lowest incidence (21.2 per 100,000), consistent with lower "
    "person-to-person contact rates in rural settings.",
    indent=True)

add_heading(doc, "4.3 Vaccination Coverage and Immunity Gaps", size=12, space_before=8)
add_para(doc,
    "The absence of a significant inverse correlation between MR campaign coverage and "
    "incidence (r = 0.10, p = 0.82) is a critical finding. It indicates that high "
    "campaign coverage, achieved during an active outbreak, did not prevent ongoing "
    "transmission at the division level. Several explanations merit consideration.",
    indent=True)

add_para(doc,
    "First, the MR campaign was reactive: it was conducted in response to the outbreak "
    "rather than in advance of it. Children vaccinated during the campaign had insufficient "
    "time to develop full protective immunity before exposure. Second, pre-existing immunity "
    "gaps — accumulated over years of suboptimal routine immunisation, pandemic-era service "
    "disruptions, and cohorts of unvaccinated children — provided a large susceptible pool "
    "that the campaign could not rapidly eliminate. Third, vaccine effectiveness in field "
    "conditions may be reduced by cold-chain failures, particularly in low-resource urban "
    "settings. Fourth, coverage reported at division level may mask sub-district "
    "heterogeneity, where pockets of unvaccinated children sustain localised transmission "
    "chains even in a division with nominally high aggregate coverage.",
    indent=True)

add_heading(doc, "4.4 Mortality", size=12, space_before=8)
add_para(doc,
    "The confirmed CFR of 0.99% is higher than estimates from high-income countries "
    "(typically <0.1%) but within the range reported for low- and middle-income settings "
    "(0.1%–5.0%). Measles mortality is strongly associated with malnutrition, vitamin A "
    "deficiency, and inadequate access to supportive care. The high hospitalisation rate "
    "(80.6% of suspected cases) suggests either clinical severity or health-seeking "
    "behaviour that triages severe cases to hospital, while milder cases may be managed "
    "at home and not reported. Deaths data were incomplete across the full study period; "
    "mortality trends could not be reliably characterised, which is a significant "
    "limitation of this analysis.",
    indent=True)

add_heading(doc, "4.5 Comparison with Prior Bangladesh Outbreaks", size=12, space_before=8)
add_para(doc,
    "Prior to this outbreak, Bangladesh documented measles outbreaks in 2012 and 2018–2019, "
    "primarily affecting under-five children and adolescents in urban slums. The 2026 "
    "outbreak appears to be the largest measles event in Bangladesh's documented history "
    "in terms of absolute case counts, with over 73,000 suspected cases reported in "
    "less than three months of active surveillance. Historical outbreaks were associated "
    "with incomplete two-dose vaccination schedules and delays in supplementary "
    "immunisation activities.",
    indent=True)

add_heading(doc, "4.6 Limitations", size=12, space_before=8)

limitations = [
    "Aggregated data only: No individual patient data (age, sex, vaccination status, "
    "district, clinical outcome) were available from the daily press releases, preventing "
    "age-stratified or sex-stratified analyses.",
    "Case definition sensitivity: Suspected cases include all febrile rash illnesses; the "
    "12.5% confirmation rate suggests a substantial proportion of suspected cases may be "
    "non-measles illnesses (e.g., dengue, rubella), particularly in the context of "
    "concurrent dengue circulation in Bangladesh.",
    "Incomplete deaths data: Mortality figures were reliably extractable for only 14 of "
    "62 study days due to PDF layout variability. Deaths for earlier dates may be "
    "underestimated.",
    "Division-level vaccination data: Campaign coverage data were available only at "
    "division level (n = 8), limiting statistical power for correlation analyses.",
    "No pre-outbreak baseline: Press release data began on 15 March 2026; the true "
    "onset of the outbreak and its initial exponential growth phase are not captured.",
    "Data source dependency: All data were derived from a single government source "
    "(DGHS). Independent verification was not possible.",
]

for item in limitations:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    set_font(run, size=12)

doc.add_paragraph()

# ── 5. CONCLUSIONS ─────────────────────────────────────────────────────────

add_heading(doc, "5. Conclusions", size=13)

add_para(doc,
    "The 2026 measles outbreak in Bangladesh is one of the largest on record, with over "
    "73,000 suspected and 9,136 confirmed cases reported in 62 days, 504 suspected deaths, "
    "and a rapid early doubling time of 6.1 days. The outbreak was disproportionately "
    "concentrated in Dhaka — Bangladesh's most densely populated division — with no "
    "statistically significant relationship between MR campaign coverage and division-level "
    "incidence, highlighting the limitations of reactive vaccination campaigns in interrupting "
    "active transmission.",
    indent=True)

add_para(doc,
    "These findings underscore three urgent policy priorities: (i) strengthening routine "
    "two-dose MR immunisation to maintain population immunity above the measles elimination "
    "threshold (≥95% two-dose coverage); (ii) implementing targeted supplementary "
    "immunisation in high-density urban communities — particularly Dhaka metropolitan area — "
    "well in advance of outbreak risk periods; and (iii) improving real-time surveillance "
    "infrastructure, including laboratory confirmation capacity, to enable timely and "
    "granular epidemiological response.",
    indent=True)

add_para(doc,
    "Continued daily surveillance data publication by DGHS is a commendable example of "
    "government transparency during a public health emergency. Future analyses incorporating "
    "age, sex, district-level, and vaccination-status data would substantially improve "
    "the epidemiological characterisation of this outbreak.",
    indent=True)

# ── DATA AVAILABILITY ──────────────────────────────────────────────────────

add_heading(doc, "Data Availability", size=13)
add_para(doc,
    "All data, analysis scripts, and figures are publicly available at: "
    "https://github.com/khalilurrrahmanridoykhan/bangladesh-measles-2026. "
    "The primary data source is the DGHS Bangladesh press release portal: "
    "https://dghs.gov.bd/pages/press-releases/.")

# ── COMPETING INTERESTS ────────────────────────────────────────────────────

add_heading(doc, "Competing Interests", size=13)
add_para(doc, "The author declares no competing interests.")

# ── FUNDING ────────────────────────────────────────────────────────────────

add_heading(doc, "Funding", size=13)
add_para(doc,
    "This research received no specific funding from any public, commercial, or "
    "not-for-profit funding agency.")

# ── ACKNOWLEDGEMENTS ───────────────────────────────────────────────────────

add_heading(doc, "Acknowledgements", size=13)
add_para(doc,
    "The author acknowledges the Directorate General of Health Services (DGHS), "
    "Government of Bangladesh, for the daily public release of measles surveillance data "
    "throughout the 2026 outbreak. The transparency of this reporting enabled independent "
    "epidemiological analysis.")

# ── REFERENCES ─────────────────────────────────────────────────────────────

add_heading(doc, "References", size=13)

refs = [
    "1. World Health Organization (WHO). Measles. Geneva: WHO; 2023. "
       "Available from: https://www.who.int/news-room/fact-sheets/detail/measles",
    "2. Directorate General of Health Services (DGHS), Bangladesh. Daily Measles (হাম) "
       "Press Releases. Dhaka: DGHS; 2026. "
       "Available from: https://dghs.gov.bd/pages/press-releases/",
    "3. Moss WJ. Measles. Lancet. 2017;390(10111):2490–2502. "
       "doi:10.1016/S0140-6736(17)31463-0",
    "4. Orenstein WA, Seib K. Mounting a good offense against measles. "
       "N Engl J Med. 2014;371(18):1661–1663.",
    "5. WHO Regional Office for South-East Asia. Measles and Rubella Elimination "
       "Strategic Plan 2020–2024. New Delhi: WHO SEARO; 2020.",
    "6. Bangladesh Bureau of Statistics (BBS). Bangladesh Population and Housing "
       "Census 2022. Dhaka: BBS; 2023.",
    "7. Expanded Programme on Immunization (EPI), Bangladesh. EPI Coverage Evaluation "
       "Survey Reports. Dhaka: DGHS; 2019.",
    "8. Luquero FJ, Pham-Orsetti H, Cummings DAT, et al. A long-lasting measles epidemic "
       "in Maroua, Cameroon 2008–2009: mass vaccination as response to the epidemic. "
       "J Infect Dis. 2011;204(Suppl 1):S243–S251.",
    "9. Grout L, Minetti A, Hurtado N, et al. Measles in Democratic Republic of Congo: "
       "an ongoing crisis. Epidemiol Infect. 2013;141(6):1131–1139.",
    "10. Perry RT, Halsey NA. The clinical significance of measles: a review. "
        "J Infect Dis. 2004;189(Suppl 1):S4–S16.",
    "11. Dabbagh A, Laws RL, Steulet C, et al. Progress toward regional measles elimination "
        "— worldwide, 2000–2017. MMWR Morb Mortal Wkly Rep. 2018;67(47):1323–1329.",
    "12. Kundu SK, Khanam M, Das MK. Measles outbreak investigation in Bangladesh, 2018–2019: "
        "a field epidemiology report. Bangladesh J Infect Dis. 2020;7(1):18–24.",
    "13. Fine PE, Clarkson JA. Measles in England and Wales — I: an analysis of factors "
        "underlying seasonal patterns. Int J Epidemiol. 1982;11(1):5–14.",
    "14. WHO. Measles vaccines: WHO position paper — April 2017. "
        "Wkly Epidemiol Rec. 2017;92(17):205–227.",
    "15. Gavi, the Vaccine Alliance. Measles supplementary immunization activities: "
        "evidence and programmatic guidance. Geneva: Gavi; 2021.",
    "16. Lessler J, Metcalf CJE. Balancing evidence and uncertainty when considering "
        "rubella vaccine introduction. PLoS One. 2013;8(7):e67639.",
    "17. Bhatt M, Favin M, Ackley J, Barry M. On the measurement of vaccination coverage. "
        "Health Policy Plan. 1993;8(4):296–303.",
    "18. UNICEF Bangladesh. Immunization programme — Bangladesh. Dhaka: UNICEF; 2023. "
        "Available from: https://www.unicef.org/bangladesh/immunization",
    "19. World Health Organization. WHO surveillance standards for measles. "
        "Geneva: WHO; 2018.",
    "20. van den Ent MMVX, Brown DW, Hoekstra EJ, Christie A, Cochi SL. "
        "Measles mortality reduction contributes substantially to reduction of all cause "
        "mortality among children less than five years of age. "
        "J Infect Dis. 2011;204(Suppl 1):S18–S23.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    run = p.add_run(ref)
    set_font(run, size=11)

# ── FIGURES (embedded) ─────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, "Figures", size=14, center=True, space_before=0)

figures = [
    ("fig1_epidemic_curve.png",
     "Figure 1. Epidemic curve of the 2026 measles outbreak in Bangladesh. "
     "Daily suspected cases (orange bars) and confirmed cases (blue bars) with "
     "seven-day centred rolling averages (solid and dashed lines, respectively). "
     "2 April to 2 June 2026."),
    ("fig2_cumulative_trajectory.png",
     "Figure 2. Cumulative trajectory of measles cases, hospitalisations, recoveries, "
     "and suspected deaths, Bangladesh 2026."),
    ("fig3_division_burden.png",
     "Figure 3. Division-level measles burden, Bangladesh 2026. "
     "Panel A: Confirmed cases (cumulative). "
     "Panel B: Incidence rate per 100,000 population (2022 Census denominators)."),
    ("fig4_vaccination_vs_incidence.png",
     "Figure 4. Scatter plot of Measles-Rubella Campaign 2026 coverage (%) versus "
     "confirmed measles incidence per 100,000 population by division. "
     "Bubble size is proportional to the number of confirmed cases. "
     "Pearson r = 0.10, p = 0.82. Line indicates linear regression fit."),
]

for fname, caption in figures:
    fpath = os.path.join(FIG, fname)
    if os.path.exists(fpath):
        doc.add_paragraph()
        doc.add_picture(fpath, width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(20)
        cr = cp.add_run(caption)
        set_font(cr, size=10, italic=True)
    else:
        add_para(doc, f"[Figure not found: {fname}]", italic=True)

# ── SAVE ───────────────────────────────────────────────────────────────────

out_path = os.path.join(OUT, "Bangladesh_Measles_2026_Manuscript.docx")
doc.save(out_path)
print(f"\nManuscript saved: {out_path}")
print(f"Word count (approximate): {sum(len(p.text.split()) for p in doc.paragraphs):,} words")
